# Imports básicos e do CustomTkinter
import tkinter as tk
import customtkinter as ctk
from tkinter import messagebox

# Imports necessários para GeradorIndeterminadoFrame (do script original)
import re
from datetime import datetime
import requests
import json
from num2words import num2words
from docx import Document

# --- Placeholder Frames (Mantidos para os outros botões) ---

class GeradorDeterminadoFrame(ctk.CTkFrame):
    def __init__(self, master, controller):
        super().__init__(master)
        self.controller = controller
        self.grid_columnconfigure(0, weight=1)
        label = ctk.CTkLabel(self, text="Formulário: Contrato Prazo Determinado", font=ctk.CTkFont(size=16, weight="bold"))
        label.grid(row=0, column=0, padx=20, pady=20)
        placeholder_content = ctk.CTkLabel(self, text="Conteúdo do Gerador Determinado\n(a ser implementado)")
        placeholder_content.grid(row=1, column=0, padx=20, pady=10)
        voltar_button = ctk.CTkButton(self, text="Voltar", command=lambda: controller.show_frame("SelecaoFrame"))
        voltar_button.grid(row=99, column=0, padx=20, pady=20, sticky="s")

class TermoDividaFrame(ctk.CTkFrame):
    def __init__(self, master, controller):
        super().__init__(master)
        self.controller = controller
        self.grid_columnconfigure(0, weight=1)
        label = ctk.CTkLabel(self, text="Formulário: Termo de Dívida", font=ctk.CTkFont(size=16, weight="bold"))
        label.grid(row=0, column=0, padx=20, pady=20)
        placeholder_content = ctk.CTkLabel(self, text="Conteúdo do Termo de Dívida\n(a ser implementado)")
        placeholder_content.grid(row=1, column=0, padx=20, pady=10)
        voltar_button = ctk.CTkButton(self, text="Voltar", command=lambda: controller.show_frame("SelecaoFrame"))
        voltar_button.grid(row=99, column=0, padx=20, pady=20, sticky="s")

class InvestidorFrame(ctk.CTkFrame):
    def __init__(self, master, controller):
        super().__init__(master)
        self.controller = controller
        self.grid_columnconfigure(0, weight=1)
        label = ctk.CTkLabel(self, text="Formulário: Contrato de Investidor", font=ctk.CTkFont(size=16, weight="bold"))
        label.grid(row=0, column=0, padx=20, pady=20)
        placeholder_content = ctk.CTkLabel(self, text="Conteúdo do Contrato Investidor\n(a ser implementado)")
        placeholder_content.grid(row=1, column=0, padx=20, pady=10)
        voltar_button = ctk.CTkButton(self, text="Voltar", command=lambda: controller.show_frame("SelecaoFrame"))
        voltar_button.grid(row=99, column=0, padx=20, pady=20, sticky="s")

# --- Frame do Gerador Indeterminado (COM SCROLL) ---

class GeradorIndeterminadoFrame(ctk.CTkFrame):
    def __init__(self, master, controller):
        super().__init__(master)
        self.controller = controller
        self.veiculos_data = {}
        self.veiculo_keys = []
        self._load_vehicle_data()

        self.qtd_parcelas_caucao = ""
        self.valor_parcela_caucao = 0.0
        self.valor_extenso_caucao = ""

        # --- Variáveis dos Widgets ---
        self.nome_var = ctk.StringVar()
        self.cpf_var = ctk.StringVar()
        self.veiculo_var = ctk.StringVar(value="Selecione...")
        self.renavam_var = ctk.StringVar()
        self.chassi_var = ctk.StringVar()
        self.proprietario_var = ctk.StringVar()
        self.cpf_cnpj_proprietario_var = ctk.StringVar()
        self.cep_var = ctk.StringVar()
        self.endereco_var = ctk.StringVar()
        self.numero_var = ctk.StringVar()
        self.valor_var = ctk.StringVar()
        self.cautao_var = ctk.StringVar(value="Quitado")
        self.data_var = ctk.StringVar()
        self.email_var = ctk.StringVar()
        self.telefone_var = ctk.StringVar()

        # --- Configuração do Grid PRINCIPAL do GeradorIndeterminadoFrame ---
        self.grid_rowconfigure(1, weight=1)  # Linha do scrollable frame deve expandir
        self.grid_columnconfigure(0, weight=1) # Coluna única para expandir

        # --- Cabeçalho (Fora do Scroll) ---
        header = ctk.CTkLabel(self, text="GERADOR DE CONTRATO - PRAZO INDETERMINADO", font=ctk.CTkFont(size=16, weight="bold"))
        header.grid(row=0, column=0, padx=20, pady=(10, 15))

        # --- Frame Rolável para o Conteúdo do Formulário ---
        scrollable_content_frame = ctk.CTkScrollableFrame(self, label_text=None)
        scrollable_content_frame.grid(row=1, column=0, padx=10, pady=0, sticky="nsew")

        # Configurar grid DENTRO do frame rolável
        scrollable_content_frame.grid_columnconfigure(0, weight=1) # Coluna dos Labels
        scrollable_content_frame.grid_columnconfigure(1, weight=3) # Coluna dos Entries/Combos
        scrollable_content_frame.grid_columnconfigure(2, weight=1) # Coluna do botão Buscar CEP

        # --- Widgets do Formulário DENTRO do Frame Rolável ---
        ctk.CTkLabel(scrollable_content_frame, text="Nome:").grid(row=1, column=0, sticky="w", padx=10, pady=5)
        nome_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.nome_var)
        nome_entry.grid(row=1, column=1, columnspan=2, sticky="ew", padx=10, pady=5)

        ctk.CTkLabel(scrollable_content_frame, text="CPF:").grid(row=2, column=0, sticky="w", padx=10, pady=5)
        cpf_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.cpf_var)
        cpf_entry.grid(row=2, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        cpf_entry.bind("<FocusOut>", self._format_cpf)

        ctk.CTkLabel(scrollable_content_frame, text="Veículo:").grid(row=3, column=0, sticky="w", padx=10, pady=5)
        veiculo_cb = ctk.CTkComboBox(scrollable_content_frame, variable=self.veiculo_var, values=self.veiculo_keys,
                                     state="readonly", command=self._atualizar_dados_veiculo)
        veiculo_cb.grid(row=3, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        veiculo_cb.set("Selecione...")

        ctk.CTkLabel(scrollable_content_frame, text="Renavam:").grid(row=4, column=0, sticky="w", padx=10, pady=5)
        renavam_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.renavam_var, state="readonly")
        renavam_entry.grid(row=4, column=1, columnspan=2, sticky="ew", padx=10, pady=5)

        ctk.CTkLabel(scrollable_content_frame, text="Chassi:").grid(row=5, column=0, sticky="w", padx=10, pady=5)
        chassi_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.chassi_var, state="readonly")
        chassi_entry.grid(row=5, column=1, columnspan=2, sticky="ew", padx=10, pady=5)

        ctk.CTkLabel(scrollable_content_frame, text="Proprietário:").grid(row=6, column=0, sticky="w", padx=10, pady=5)
        proprietario_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.proprietario_var, state="readonly")
        proprietario_entry.grid(row=6, column=1, columnspan=2, sticky="ew", padx=10, pady=5)

        ctk.CTkLabel(scrollable_content_frame, text="CPF/CNPJ Proprietário:").grid(row=7, column=0, sticky="w", padx=10, pady=5)
        cpf_cnpj_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.cpf_cnpj_proprietario_var, state="readonly")
        cpf_cnpj_entry.grid(row=7, column=1, columnspan=2, sticky="ew", padx=10, pady=5)

        ctk.CTkLabel(scrollable_content_frame, text="CEP:").grid(row=8, column=0, sticky="w", padx=10, pady=5)
        cep_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.cep_var)
        cep_entry.grid(row=8, column=1, sticky="ew", padx=10, pady=5)
        btn_cep = ctk.CTkButton(scrollable_content_frame, text="Buscar CEP", width=100, command=self.buscar_cep)
        btn_cep.grid(row=8, column=2, padx=(5, 10), pady=5)

        ctk.CTkLabel(scrollable_content_frame, text="Endereço:").grid(row=9, column=0, sticky="w", padx=10, pady=5)
        endereco_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.endereco_var, state="readonly")
        endereco_entry.grid(row=9, column=1, columnspan=2, sticky="ew", padx=10, pady=5)

        ctk.CTkLabel(scrollable_content_frame, text="Número/Complemento:").grid(row=10, column=0, sticky="w", padx=10, pady=5)
        numero_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.numero_var)
        numero_entry.grid(row=10, column=1, columnspan=2, sticky="ew", padx=10, pady=5)

        ctk.CTkLabel(scrollable_content_frame, text="Email:").grid(row=11, column=0, sticky="w", padx=10, pady=5)
        email_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.email_var)
        email_entry.grid(row=11, column=1, columnspan=2, sticky="ew", padx=10, pady=5)

        ctk.CTkLabel(scrollable_content_frame, text="Telefone:").grid(row=12, column=0, sticky="w", padx=10, pady=5)
        telefone_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.telefone_var)
        telefone_entry.grid(row=12, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        telefone_entry.bind("<FocusOut>", self._format_telefone)

        ctk.CTkLabel(scrollable_content_frame, text="Valor (R$):").grid(row=13, column=0, sticky="w", padx=10, pady=5)
        valor_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.valor_var)
        valor_entry.grid(row=13, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        valor_entry.bind("<FocusOut>", self._format_valor)

        ctk.CTkLabel(scrollable_content_frame, text="Caução:").grid(row=14, column=0, sticky="w", padx=10, pady=5)
        cautao_cb = ctk.CTkComboBox(scrollable_content_frame, variable=self.cautao_var, values=["Quitado", "Parcelado"],
                                     state="readonly", command=self._on_caucao_change)
        cautao_cb.grid(row=14, column=1, columnspan=2, sticky="ew", padx=10, pady=5)

        ctk.CTkLabel(scrollable_content_frame, text="Data (dd/mm/aaaa):").grid(row=15, column=0, sticky="w", padx=10, pady=5)
        data_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.data_var, placeholder_text="dd/mm/aaaa")
        data_entry.grid(row=15, column=1, columnspan=2, sticky="ew", padx=10, pady=5)


        # --- Botões de Ação (Fora do Scroll) ---
        button_frame = ctk.CTkFrame(self, fg_color="transparent")
        button_frame.grid(row=2, column=0, pady=15)

        btn_gerar = ctk.CTkButton(button_frame, text="GERAR CONTRATO", command=self.gerar_contrato, width=200, height=40)
        btn_gerar.pack(side=tk.LEFT, padx=10)

        btn_voltar = ctk.CTkButton(button_frame, text="Voltar", command=lambda: self.controller.show_frame("SelecaoFrame"), width=100)
        btn_voltar.pack(side=tk.LEFT, padx=10)

    # --- Métodos Auxiliares ---
    def _somente_letras(self, texto):
        return re.fullmatch(r'[A-Za-zÀ-ÖØ-öø-ÿ\s]+', texto) is not None

    def _format_cpf(self, event=None):
        s = self.cpf_var.get()
        digits = re.sub(r'\D', '', s)[:11]
        if len(digits) <= 3: formatted = digits
        elif len(digits) <= 6: formatted = f"{digits[:3]}.{digits[3:]}"
        elif len(digits) <= 9: formatted = f"{digits[:3]}.{digits[3:6]}.{digits[6:]}"
        else: formatted = f"{digits[:3]}.{digits[3:6]}.{digits[6:9]}-{digits[9:]}"
        self.cpf_var.set(formatted)

    def _format_valor(self, event=None):
        s = self.valor_var.get()
        digits = re.sub(r'\D', '', s)
        if not digits: formatted = ""
        else:
            try:
                num = float(digits) / 100
                formatted = f"{num:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            except ValueError: formatted = ""
        self.valor_var.set(formatted)

    def _format_telefone(self, event=None):
        s = self.telefone_var.get()
        digits = re.sub(r'\D', '', s)[:11]
        formatted = digits
        if len(digits) == 11: formatted = f"({digits[:2]}) {digits[2]} {digits[3:7]}-{digits[7:]}"
        elif len(digits) == 10: formatted = f"({digits[:2]}) {digits[2:6]}-{digits[6:]}"
        self.telefone_var.set(formatted)

    def _load_vehicle_data(self):
        try:
            # !! IMPORTANTE: Verifique se este caminho está correto !!
            with open("json/veiculos.json", "r", encoding='utf-8') as f:
                self.veiculos_data = json.load(f)
                self.veiculo_keys = list(self.veiculos_data.keys()) if self.veiculos_data else []
        except FileNotFoundError:
             messagebox.showerror("Erro", "Arquivo 'json/veiculos.json' não encontrado.\nVerifique o caminho no código.", parent=self)
             self.veiculos_data = {}
             self.veiculo_keys = []
        except json.JSONDecodeError:
             messagebox.showerror("Erro", "Erro ao decodificar 'json/veiculos.json'.\nVerifique o formato do arquivo.", parent=self)
             self.veiculos_data = {}
             self.veiculo_keys = []
        except Exception as e:
            messagebox.showerror("Erro", f"Erro inesperado ao carregar 'veiculos.json':\n{e}", parent=self)
            self.veiculos_data = {}
            self.veiculo_keys = []
        # Atualiza o combobox caso os dados sejam carregados depois da inicialização (embora aqui seja no init)
        # Acessa o widget diretamente se necessário, ou confia no CTk para pegar os values atualizados.
        # A forma mais segura seria recriar ou reconfigurar o combobox aqui, mas vamos manter simples por ora.

    def _atualizar_dados_veiculo(self, selected_key=None):
        key = self.veiculo_var.get()
        if key in self.veiculos_data:
            data = self.veiculos_data[key]
             # Adiciona verificação se data é lista e tem elementos suficientes
            if isinstance(data, list):
                self.renavam_var.set(data[0] if len(data) > 0 else "")
                self.chassi_var.set(data[1] if len(data) > 1 else "")
                self.proprietario_var.set(data[2] if len(data) > 2 else "")
                self.cpf_cnpj_proprietario_var.set(data[3] if len(data) > 3 else "")
            else: # Limpa se o formato dos dados no JSON estiver incorreto para a chave
                 self.renavam_var.set("")
                 self.chassi_var.set("")
                 self.proprietario_var.set("")
                 self.cpf_cnpj_proprietario_var.set("")
        else:
            self.renavam_var.set("")
            self.chassi_var.set("")
            self.proprietario_var.set("")
            self.cpf_cnpj_proprietario_var.set("")

    def _abrir_configuracao_caucao(self):
        top = ctk.CTkToplevel(self)
        top.title("Configurar Parcelamento")
        top.geometry("380x200")
        top.transient(self)
        top.grab_set() # Foca nesta janela

        top.grid_columnconfigure(0, weight=1)
        top.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(top, text="Valor de Entrada (R$):").grid(row=0, column=0, padx=10, pady=10, sticky="w")
        entrada_var = ctk.StringVar()
        entrada_entry = ctk.CTkEntry(top, textvariable=entrada_var, placeholder_text="Ex: 300,00")
        entrada_entry.grid(row=0, column=1, padx=10, pady=10, sticky="ew")
        # Função interna para formatação, para não poluir a classe principal
        def _format_entrada_toplevel(event=None):
            s = entrada_var.get(); digits = re.sub(r'\D', '', s)
            if not digits: formatted = ""
            else:
                try: num = float(digits)/100; formatted = f"{num:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                except ValueError: formatted = ""
            entrada_var.set(formatted)
        entrada_entry.bind("<FocusOut>", _format_entrada_toplevel)

        ctk.CTkLabel(top, text="Quantidade de Semanas:").grid(row=1, column=0, padx=10, pady=10, sticky="w")
        semanas_var = ctk.StringVar()
        semanas_entry = ctk.CTkEntry(top, textvariable=semanas_var)
        semanas_entry.grid(row=1, column=1, padx=10, pady=10, sticky="ew")

        def confirmar_parcelamento():
            try:
                entrada_str = entrada_var.get().replace(".", "").replace(",", ".")
                entrada = float(entrada_str) if entrada_str else 0.0
                semanas_str = semanas_var.get()
                # Garante que semanas seja um inteiro positivo
                semanas = int(semanas_str) if semanas_str.isdigit() and int(semanas_str) > 0 else 0

                # Define um valor padrão razoável ou pegue de uma config
                valor_total_caucao = 1000.0

                if not (0 <= entrada <= valor_total_caucao):
                    messagebox.showerror("Erro", f"Valor de entrada deve ser entre 0,00 e {valor_total_caucao:,.2f}.", parent=top); return
                if semanas < 1:
                    messagebox.showerror("Erro", "Quantidade de semanas inválida ou menor que 1.", parent=top); return

                restante = valor_total_caucao - entrada
                # Evita divisão por zero caso semanas seja 0 (embora validado acima)
                parcela = round(restante / semanas, 2) if semanas > 0 else 0.0

                self.qtd_parcelas_caucao = str(semanas)
                self.valor_parcela_caucao = parcela
                # Usar to='currency' para obter "reais" e "centavos"
                self.valor_extenso_caucao = num2words(parcela, lang='pt_BR', to='currency').upper()
                top.destroy()
            except ValueError: messagebox.showerror("Erro", "Valores de entrada ou semanas inválidos. Use números.", parent=top)
            except Exception as e: messagebox.showerror("Erro", f"Ocorreu um erro inesperado: {e}", parent=top)

        btn_confirma = ctk.CTkButton(top, text="Confirmar", command=confirmar_parcelamento)
        btn_confirma.grid(row=2, column=0, columnspan=2, pady=20)
        # Garante que a janela Toplevel apareça na frente e foca no primeiro campo
        top.after(100, top.lift)
        entrada_entry.focus()

    def _on_caucao_change(self, choice):
        if choice == "Parcelado":
            self._abrir_configuracao_caucao()
            # Se o usuário cancelar a janela toplevel, os valores podem não ser setados.
            # Adicionar uma checagem em gerar_contrato se cautao=="Parcelado"
            # e os valores self.qtd_parcelas_caucao etc. estão vazios/inválidos.
        else:
            self.qtd_parcelas_caucao = ""
            self.valor_parcela_caucao = 0.0
            self.valor_extenso_caucao = ""

    # --- Métodos de Ação ---
    def buscar_cep(self):
        cep = self.cep_var.get().strip().replace('-', '').replace('.', '')
        if not re.fullmatch(r'\d{8}', cep):
            messagebox.showerror("Erro", "CEP inválido. Informe 8 dígitos numéricos.", parent=self); return
        try:
            self.controller.update_statusbar("Buscando CEP...") # Feedback na status bar
            response = requests.get(f"https://viacep.com.br/ws/{cep}/json/", timeout=10) # Timeout
            response.raise_for_status()
            data = response.json()
            if data.get("erro"):
                messagebox.showerror("Erro", "CEP não encontrado.", parent=self)
                self.endereco_var.set("")
            else:
                logradouro = data.get("logradouro", ""); bairro = data.get("bairro", "")
                localidade = data.get("localidade", ""); uf = data.get("uf", "")
                endereco_completo = f"{logradouro}, {bairro}, {localidade} - {uf}".strip(", ")
                self.endereco_var.set(endereco_completo)
            self.controller.update_statusbar("Pronto") # Volta status normal
        except requests.exceptions.Timeout:
            messagebox.showerror("Erro de Conexão", "Tempo limite excedido ao buscar CEP.", parent=self)
            self.controller.update_statusbar("Erro ao buscar CEP")
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Erro de Conexão", f"Não foi possível conectar ao ViaCEP:\n{e}", parent=self)
            self.controller.update_statusbar("Erro ao buscar CEP")
        except json.JSONDecodeError:
            messagebox.showerror("Erro", "Resposta inválida do ViaCEP.", parent=self)
            self.controller.update_statusbar("Erro ao buscar CEP")
        except Exception as e:
            messagebox.showerror("Erro Inesperado", f"Ocorreu um erro:\n{e}", parent=self)
            self.controller.update_statusbar("Erro")


    def gerar_contrato(self):
        # Coleta
        nome=self.nome_var.get().strip(); cpf=self.cpf_var.get().strip(); veiculo=self.veiculo_var.get().strip()
        renavam=self.renavam_var.get().strip(); chassi=self.chassi_var.get().strip(); proprietario=self.proprietario_var.get().strip()
        cpf_cnpj_proprietario=self.cpf_cnpj_proprietario_var.get().strip(); cep=self.cep_var.get().strip()
        endereco=self.endereco_var.get().strip(); numero_comp=self.numero_var.get().strip()
        valor_str=self.valor_var.get().strip(); cautao=self.cautao_var.get().strip()
        data_str=self.data_var.get().strip(); email=self.email_var.get().strip(); telefone=self.telefone_var.get().strip()

        # Validação
        errors = []
        if not nome: errors.append("Nome é obrigatório.")
        elif not self._somente_letras(nome): errors.append("Nome deve conter apenas letras/espaços.")
        if not re.fullmatch(r'\d{3}\.\d{3}\.\d{3}-\d{2}', cpf): errors.append("CPF inválido (000.000.000-00).")
        if not veiculo or veiculo == "Selecione...": errors.append("Selecione um veículo.")
        if not renavam: errors.append("Renavam não preenchido.")
        if not chassi: errors.append("Chassi não preenchido.")
        if not proprietario: errors.append("Proprietário não preenchido.")
        if not cpf_cnpj_proprietario: errors.append("CPF/CNPJ Proprietário não preenchido.")
        if not cep: errors.append("CEP obrigatório.")
        if not endereco: errors.append("Busque o CEP.")
        if not numero_comp: errors.append("Número/Complemento obrigatório.")
        if not valor_str or not re.fullmatch(r'(\d{1,3}(\.\d{3})*|\d+),\d{2}', valor_str): errors.append("Valor inválido (ex: 100,00).")
        if not email or not re.fullmatch(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email): errors.append("Email inválido.")
        if not telefone or not re.fullmatch(r'\(\d{2}\)\s\d\s\d{4}-\d{4}', telefone):
             digits_tel = re.sub(r'\D', '', telefone)
             if len(digits_tel) != 11: errors.append("Telefone inválido ((XX) X XXXX-XXXX).")
             else: telefone = f"({digits_tel[:2]}) {digits_tel[2]} {digits_tel[3:7]}-{digits_tel[7:]}"; self.telefone_var.set(telefone) # Auto-formata
        try: datetime.strptime(data_str, "%d/%m/%Y")
        except ValueError: errors.append("Data inválida (dd/mm/aaaa).")
        if cautao == "Parcelado" and (not self.qtd_parcelas_caucao or self.valor_parcela_caucao <= 0):
             errors.append("Configure corretamente o parcelamento da caução.")

        if errors:
            messagebox.showerror("Erros de Validação", "\n".join(errors), parent=self)
            return

        # Preparação dos Dados
        nome_upper=nome.upper(); endereco_final=f"{endereco}, {numero_comp}"; fabricacao_modelo,placa="",""
        if " - " in veiculo:
            try: fabricacao_modelo, placa = veiculo.split(" - ", 1)
            except ValueError: fabricacao_modelo = veiculo
        else: fabricacao_modelo = veiculo
        try:
             valor_num=float(valor_str.replace(".","").replace(",","."))
             valor_extenso=num2words(valor_num,lang='pt_BR',to='currency').upper()
        except Exception as e:
             print(f"Erro converter valor extenso: {e}")
             messagebox.showwarning("Aviso", "Não foi possível converter o valor para extenso.", parent=self)
             valor_extenso=f"R$ {valor_str}" # Usa o valor numérico formatado como fallback

        qtd_parc_caucao_str=""; valor_parc_caucao_fmt=""; valor_ext_parc_caucao=""
        if cautao=="Parcelado":
            qtd_parc_caucao_str=self.qtd_parcelas_caucao
            valor_parc_caucao_fmt=f"{self.valor_parcela_caucao:,.2f}".replace(",","X").replace(".",",").replace("X",".")
            valor_ext_parc_caucao=self.valor_extenso_caucao # Já vem formatado de _abrir_config...

        # Seleção e Processamento do DOCX
        try:
            if cautao == "Parcelado": model_file = "doc/modelo_contrato_prazoinderteminadoparcelado.docx"
            else: model_file = "doc/modelo_contrato_prazoinderteminadoquitado.docx"
            # !! IMPORTANTE: Verifique se este caminho está correto !!
            doc = Document(model_file)
        except FileNotFoundError: return messagebox.showerror("Erro", f"Arquivo de modelo '{model_file}' não encontrado na pasta 'doc/'.\nVerifique o caminho.", parent=self)
        except Exception as e: return messagebox.showerror("Erro", f"Erro ao abrir o modelo DOCX:\n{e}", parent=self)

        placeholders = {
            "{{NOME}}":nome_upper, "{{CPF}}":cpf, "{{ENDERECO}}":endereco_final, "{{VEICULO}}":veiculo,
            "{{FABRICACAO_MODELO}}":fabricacao_modelo, "{{PLACA}}":placa, "{{CHASSI}}":chassi, "{{RENAVAM}}":renavam,
            "{{PROPRIETARIO}}":proprietario, "{{CPF_CNPJ_PROPRIETARIO}}":cpf_cnpj_proprietario, "{{VALOR}}":valor_str,
            "{{VALOR_EXTENSO}}":valor_extenso, "{{QTD_PARCELAS}}":qtd_parc_caucao_str,
            "{{VALOR_PARCELAS_CAUCAO}}":valor_parc_caucao_fmt, "{{VALOR_EXTENSO_CAUCAO}}":valor_ext_parc_caucao,
            "{{DATA}}":data_str, "{{EMAIL}}":email, "{{TELEFONE}}":telefone
        }

        # Substituição mais robusta (tentativa)
        try:
            for para in doc.paragraphs:
                for code, text in placeholders.items():
                    if code in para.text:
                        for run in para.runs:
                            if code in run.text:
                                # Simples replace pode quebrar formatação se o placeholder estiver dividido
                                # Uma abordagem mais segura envolveria reconstruir o run ou parágrafo
                                # Mas para casos simples, isso pode funcionar:
                                run.text = run.text.replace(code, str(text))

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                             for code, text in placeholders.items():
                                if code in para.text:
                                    for run in para.runs:
                                         if code in run.text:
                                             run.text = run.text.replace(code, str(text))

            # Salvar
            safe_nome = re.sub(r'[^\w\-\.]+', '_', nome_upper).strip('_') # Remove caracteres invalidos
            safe_veiculo = re.sub(r'[^\w\-\.]+', '_', veiculo).strip('_')
            safe_data = data_str.replace("/", "-")
            timestamp = datetime.now().strftime("%H%M%S") # Adiciona timestamp para evitar sobrescrever se gerar rápido
            novo_nome = f"Contrato_{safe_nome}_{safe_veiculo}_{safe_data}_{timestamp}.docx"

            doc.save(novo_nome)
            messagebox.showinfo("Sucesso", f"Contrato gerado com sucesso!\n\nArquivo salvo como:\n{novo_nome}", parent=self)

        except Exception as e:
             messagebox.showerror("Erro", f"Erro ao preencher ou salvar o contrato DOCX:\n{e}", parent=self)

# --- Frame do Gerador Determinado (AGORA COM CONTEÚDO E SCROLL) ---

class GeradorDeterminadoFrame(ctk.CTkFrame):
    def __init__(self, master, controller):
        super().__init__(master)
        self.controller = controller
        self.veiculos_data = {}
        self.veiculo_keys = []
        self._load_vehicle_data() # Carrega dados dos veículos

        # --- Variáveis dos Widgets ---
        self.nome_var = ctk.StringVar()
        self.cpf_var = ctk.StringVar()
        self.cep_var = ctk.StringVar()
        self.endereco_var = ctk.StringVar()
        self.numero_var = ctk.StringVar()
        self.email_var = ctk.StringVar()
        self.telefone_var = ctk.StringVar()
        self.veiculo_var = ctk.StringVar(value="Selecione...")
        self.chassi_var = ctk.StringVar()
        self.renavam_var = ctk.StringVar()
        self.proprietario_var = ctk.StringVar()
        self.cpf_cnpj_proprietario_var = ctk.StringVar()
        self.periodo_var = ctk.StringVar()
        self.valor_var = ctk.StringVar() # Valor da diária/unidade
        self.discount_var = ctk.StringVar(value="0") # Desconto percentual, inicia com 0
        self.data_var = ctk.StringVar()

        # Variáveis para campos calculados/informativos (somente leitura)
        self.periodo_extenso_var = ctk.StringVar()
        self.valor_total_var = ctk.StringVar() # Valor * Período (antes do desconto)
        self.valor_total_extenso_var = ctk.StringVar()
        self.result_var = ctk.StringVar() # Valor final com desconto

        # --- Configuração do Grid PRINCIPAL do Frame ---
        self.grid_rowconfigure(1, weight=1)  # Linha do scrollable frame expande
        self.grid_columnconfigure(0, weight=1) # Coluna única expande

        # --- Cabeçalho ---
        header = ctk.CTkLabel(self, text="GERADOR DE CONTRATO - PRAZO DETERMINADO", font=ctk.CTkFont(size=16, weight="bold"))
        header.grid(row=0, column=0, padx=20, pady=(10, 15))

        # --- Frame Rolável ---
        scrollable_content_frame = ctk.CTkScrollableFrame(self, label_text=None)
        scrollable_content_frame.grid(row=1, column=0, padx=10, pady=0, sticky="nsew")

        # Configurar grid DENTRO do frame rolável
        scrollable_content_frame.grid_columnconfigure(0, weight=1) # Labels
        scrollable_content_frame.grid_columnconfigure(1, weight=2) # Entries
        scrollable_content_frame.grid_columnconfigure(2, weight=1) # Botão CEP / Espaço

        # --- Widgets do Formulário DENTRO do Frame Rolável ---
        current_row = 0 # Contador de linha dentro do scrollable_frame

        ctk.CTkLabel(scrollable_content_frame, text="Nome:").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        nome_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.nome_var)
        nome_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="CPF:").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        cpf_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.cpf_var)
        cpf_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        cpf_entry.bind("<FocusOut>", self._format_cpf)
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="CEP:").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        cep_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.cep_var)
        cep_entry.grid(row=current_row, column=1, sticky="ew", padx=10, pady=5)
        btn_cep = ctk.CTkButton(scrollable_content_frame, text="Buscar CEP", width=100, command=self.buscar_cep)
        btn_cep.grid(row=current_row, column=2, padx=(5, 10), pady=5)
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="Endereço:").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        endereco_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.endereco_var, state="readonly")
        endereco_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="Número/Compl.:").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        numero_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.numero_var)
        numero_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="Email:").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        email_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.email_var)
        email_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="Telefone:").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        telefone_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.telefone_var)
        telefone_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        telefone_entry.bind("<FocusOut>", self._format_telefone)
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="Veículo:").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        veiculo_cb = ctk.CTkComboBox(scrollable_content_frame, variable=self.veiculo_var, values=self.veiculo_keys,
                                     state="readonly", command=self._atualizar_dados_veiculo)
        veiculo_cb.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        veiculo_cb.set("Selecione...")
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="Chassi:").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        chassi_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.chassi_var, state="readonly")
        chassi_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="Renavam:").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        renavam_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.renavam_var, state="readonly")
        renavam_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="Proprietário:").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        proprietario_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.proprietario_var, state="readonly")
        proprietario_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="CPF/CNPJ Prop.:").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        cpf_cnpj_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.cpf_cnpj_proprietario_var, state="readonly")
        cpf_cnpj_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="Período (dias):").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        periodo_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.periodo_var)
        periodo_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        periodo_entry.bind("<FocusOut>", self._format_periodo_and_calculate) # Calcula ao sair
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="Valor Diária (R$):").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        valor_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.valor_var)
        valor_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        valor_entry.bind("<FocusOut>", self._format_valor_and_calculate) # Calcula ao sair
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="Desconto (%):").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        discount_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.discount_var)
        discount_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        discount_entry.bind("<FocusOut>", self._calculate_totals) # Recalcula ao sair
        current_row += 1

        # --- Campos Calculados (Readonly) ---
        ctk.CTkLabel(scrollable_content_frame, text="Período Extenso:").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        periodo_extenso_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.periodo_extenso_var, state="readonly")
        periodo_extenso_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="Valor Total (R$):").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        valor_total_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.valor_total_var, state="readonly")
        valor_total_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="Valor Total Extenso:").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        valor_total_extenso_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.valor_total_extenso_var, state="readonly")
        valor_total_extenso_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="Resultado (c/ Desc.):").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        result_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.result_var, state="readonly")
        result_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        current_row += 1

        ctk.CTkLabel(scrollable_content_frame, text="Data (dd/mm/aaaa):").grid(row=current_row, column=0, sticky="w", padx=10, pady=5)
        data_entry = ctk.CTkEntry(scrollable_content_frame, textvariable=self.data_var, placeholder_text="dd/mm/aaaa")
        data_entry.grid(row=current_row, column=1, columnspan=2, sticky="ew", padx=10, pady=5)
        current_row += 1


        # --- Botões de Ação (Fora do Scroll) ---
        button_frame = ctk.CTkFrame(self, fg_color="transparent")
        button_frame.grid(row=2, column=0, pady=15)

        btn_gerar = ctk.CTkButton(button_frame, text="GERAR CONTRATO", command=self.gerar_contrato, width=200, height=40)
        btn_gerar.pack(side=tk.LEFT, padx=10)

        btn_voltar = ctk.CTkButton(button_frame, text="Voltar", command=lambda: self.controller.show_frame("SelecaoFrame"), width=100)
        btn_voltar.pack(side=tk.LEFT, padx=10)


    # --- Métodos Auxiliares ---
    def _somente_letras(self, texto):
        return re.fullmatch(r'[A-Za-zÀ-ÖØ-öø-ÿ\s]+', texto) is not None

    def _format_cpf(self, event=None):
        s = self.cpf_var.get()
        digits = re.sub(r'\D', '', s)[:11]
        if len(digits) <= 3: formatted = digits
        elif len(digits) <= 6: formatted = f"{digits[:3]}.{digits[3:]}"
        elif len(digits) <= 9: formatted = f"{digits[:3]}.{digits[3:6]}.{digits[6:]}"
        else: formatted = f"{digits[:3]}.{digits[3:6]}.{digits[6:9]}-{digits[9:]}"
        self.cpf_var.set(formatted)

    def _format_valor(self, valor_str):
        """Formata string de valor para padrão moeda BRL. Retorna string formatada."""
        digits = re.sub(r'\D', '', valor_str)
        if not digits: return ""
        try:
            num = float(digits) / 100
            formatted = f"{num:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            return formatted
        except ValueError:
            return ""

    def _format_valor_and_calculate(self, event=None):
        """Formata campo VALOR DIÁRIA e recalcula totais."""
        formatted = self._format_valor(self.valor_var.get())
        self.valor_var.set(formatted)
        self._calculate_totals() # Recalcula tudo

    def _format_periodo_and_calculate(self, event=None):
        """Formata campo PERIODO, atualiza extenso e recalcula totais."""
        s = self.periodo_var.get().strip()
        periodo_extenso_txt = ""
        if s.isdigit():
            periodo_int = int(s)
            if periodo_int > 0:
                 try:
                     # Usar to='cardinal' para obter só o número por extenso
                     num_extenso = num2words(periodo_int, lang='pt_BR', to='cardinal').upper()
                     periodo_extenso_txt = f"{periodo_int} ({num_extenso}) DIAS"
                 except Exception as e:
                      print(f"Erro num2words periodo: {e}")
                      periodo_extenso_txt = f"{periodo_int} DIAS" # Fallback
            else: # Se for 0 ou negativo
                self.periodo_var.set("") # Limpa se inválido
        else: # Se não for dígito
             self.periodo_var.set("") # Limpa se inválido

        self.periodo_extenso_var.set(periodo_extenso_txt)
        self._calculate_totals() # Recalcula tudo

    def _calculate_totals(self, event=None):
        """Calcula valor total, valor final (com desconto) e atualiza os campos."""
        valor_total_num = 0.0
        resultado_num = 0.0
        valor_total_fmt = ""
        valor_total_extenso_txt = ""
        resultado_fmt = ""

        periodo_str = self.periodo_var.get().strip()
        valor_diaria_str_fmt = self.valor_var.get()
        desconto_str = self.discount_var.get().strip().replace('%','')

        try:
            periodo = int(periodo_str) if periodo_str.isdigit() else 0
            valor_diaria = float(valor_diaria_str_fmt.replace(".", "").replace(",", ".")) if valor_diaria_str_fmt else 0.0
            # Trata desconto - permite vírgula ou ponto como decimal
            desconto_str_cleaned = desconto_str.replace(",",".")
            desconto_perc = float(desconto_str_cleaned) if desconto_str_cleaned else 0.0

            if periodo > 0 and valor_diaria > 0:
                valor_total_num = valor_diaria * periodo
                valor_total_fmt = self._format_valor(str(int(valor_total_num * 100))) # Formata o total

                try:
                     valor_total_extenso_txt = num2words(valor_total_num, lang='pt_BR', to='currency').upper()
                except Exception as e:
                    print(f"Erro num2words total: {e}")
                    valor_total_extenso_txt = f"R$ {valor_total_fmt}" # Fallback

                if 0 <= desconto_perc <= 100:
                    resultado_num = valor_total_num * (1 - desconto_perc / 100)
                    resultado_fmt = self._format_valor(str(int(round(resultado_num * 100)))) # Formata resultado
                else: # Desconto inválido
                    self.discount_var.set("0") # Reseta desconto para 0
                    resultado_num = valor_total_num
                    resultado_fmt = valor_total_fmt

        except ValueError:
            # Erro na conversão de número, provavelmente formato inválido
            # Os campos já devem estar formatados ou vazios pelos binds, mas como segurança:
             pass
        except Exception as e:
             print(f"Erro inesperado em _calculate_totals: {e}")

        # Atualiza os campos readonly
        self.valor_total_var.set(valor_total_fmt)
        self.valor_total_extenso_var.set(valor_total_extenso_txt)
        self.result_var.set(resultado_fmt)


    def _format_telefone(self, event=None):
        s = self.telefone_var.get()
        digits = re.sub(r'\D', '', s)[:11]
        formatted = digits
        if len(digits) == 11: formatted = f"({digits[:2]}) {digits[2]} {digits[3:7]}-{digits[7:]}"
        elif len(digits) == 10: formatted = f"({digits[:2]}) {digits[2:6]}-{digits[6:]}"
        self.telefone_var.set(formatted)

    def _load_vehicle_data(self):
        try:
            with open("json/veiculos.json", "r", encoding='utf-8') as f:
                self.veiculos_data = json.load(f)
                self.veiculo_keys = list(self.veiculos_data.keys()) if self.veiculos_data else []
        except FileNotFoundError:
             messagebox.showerror("Erro", "Arquivo 'json/veiculos.json' não encontrado.", parent=self)
             self.veiculos_data = {}; self.veiculo_keys = []
        except json.JSONDecodeError:
             messagebox.showerror("Erro", "Erro ao decodificar 'json/veiculos.json'.", parent=self)
             self.veiculos_data = {}; self.veiculo_keys = []
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar 'veiculos.json': {e}", parent=self)
            self.veiculos_data = {}; self.veiculo_keys = []

    def _atualizar_dados_veiculo(self, selected_key=None):
        key = self.veiculo_var.get()
        if key in self.veiculos_data:
            data = self.veiculos_data[key]
            if isinstance(data, list):
                self.renavam_var.set(data[0] if len(data) > 0 else "")
                self.chassi_var.set(data[1] if len(data) > 1 else "")
                self.proprietario_var.set(data[2] if len(data) > 2 else "")
                self.cpf_cnpj_proprietario_var.set(data[3] if len(data) > 3 else "")
            else: self._limpar_dados_veiculo()
        else: self._limpar_dados_veiculo()

    def _limpar_dados_veiculo(self):
        """ Limpa os campos dependentes do veículo """
        self.renavam_var.set("")
        self.chassi_var.set("")
        self.proprietario_var.set("")
        self.cpf_cnpj_proprietario_var.set("")

    # --- Métodos de Ação ---
    def buscar_cep(self):
        cep = self.cep_var.get().strip().replace('-', '').replace('.', '')
        if not re.fullmatch(r'\d{8}', cep):
            messagebox.showerror("Erro", "CEP inválido.", parent=self); return
        try:
            self.controller.update_statusbar("Buscando CEP...")
            response = requests.get(f"https://viacep.com.br/ws/{cep}/json/", timeout=10)
            response.raise_for_status(); data = response.json()
            if data.get("erro"):
                messagebox.showerror("Erro", "CEP não encontrado.", parent=self); self.endereco_var.set("")
            else:
                logradouro=data.get("logradouro",""); bairro=data.get("bairro","")
                localidade=data.get("localidade",""); uf=data.get("uf","")
                endereco_completo = f"{logradouro}, {bairro}, {localidade} - {uf}".strip(", ")
                self.endereco_var.set(endereco_completo)
            self.controller.update_statusbar("Pronto")
        except requests.exceptions.Timeout:
            messagebox.showerror("Erro", "Tempo limite excedido ao buscar CEP.", parent=self)
            self.controller.update_statusbar("Erro CEP (Timeout)")
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Erro", f"Erro de conexão ViaCEP:\n{e}", parent=self)
            self.controller.update_statusbar("Erro CEP (Conexão)")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao buscar CEP:\n{e}", parent=self)
            self.controller.update_statusbar("Erro CEP")

    def gerar_contrato(self):
        # --- Coleta ---
        nome=self.nome_var.get().strip(); cpf=self.cpf_var.get().strip(); veiculo=self.veiculo_var.get().strip()
        renavam=self.renavam_var.get().strip(); chassi=self.chassi_var.get().strip(); proprietario=self.proprietario_var.get().strip()
        cpf_cnpj_proprietario=self.cpf_cnpj_proprietario_var.get().strip(); cep=self.cep_var.get().strip()
        endereco=self.endereco_var.get().strip(); numero_comp=self.numero_var.get().strip()
        valor_diaria_str=self.valor_var.get().strip(); # Diária
        data_str=self.data_var.get().strip(); email=self.email_var.get().strip(); telefone=self.telefone_var.get().strip()
        periodo_str = self.periodo_var.get().strip()
        discount_str_input = self.discount_var.get().strip().replace('%','') # Desconto % (input)

        # Recalcular valores finais para garantir consistência antes de usar
        self._calculate_totals()
        # Obter os valores calculados das StringVars
        valor_total_str = self.valor_total_var.get()
        valor_total_extenso = self.valor_total_extenso_var.get()
        periodo_extenso = self.periodo_extenso_var.get()
        resultado_str = self.result_var.get() # Valor final com desconto

        # --- Validação ---
        errors = []
        if not nome: errors.append("Nome é obrigatório.")
        elif not self._somente_letras(nome): errors.append("Nome inválido.")
        if not re.fullmatch(r'\d{3}\.\d{3}\.\d{3}-\d{2}', cpf): errors.append("CPF inválido (000.000.000-00).")
        if not veiculo or veiculo == "Selecione...": errors.append("Selecione um veículo.")
        if not renavam: errors.append("Renavam não preenchido.")
        if not chassi: errors.append("Chassi não preenchido.")
        if not proprietario: errors.append("Proprietário não preenchido.")
        if not cpf_cnpj_proprietario: errors.append("CPF/CNPJ Proprietário não preenchido.")
        if not cep: errors.append("CEP obrigatório.")
        if not endereco: errors.append("Busque o CEP.")
        if not numero_comp: errors.append("Número/Complemento obrigatório.")
        if not valor_diaria_str or not re.fullmatch(r'(\d{1,3}(\.\d{3})*|\d+),\d{2}', valor_diaria_str): errors.append("Valor Diária inválido (ex: 100,00).")
        if not email or not re.fullmatch(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email): errors.append("Email inválido.")
        if not telefone or not re.fullmatch(r'\(\d{2}\)\s\d\s\d{4}-\d{4}', telefone): errors.append("Telefone inválido ((XX) X XXXX-XXXX).")
        try: datetime.strptime(data_str, "%d/%m/%Y")
        except ValueError: errors.append("Data inválida (dd/mm/aaaa).")
        if not periodo_str or not periodo_str.isdigit() or int(periodo_str) <= 0: errors.append("Período (dias) inválido.")
        try: float(discount_str_input.replace(",",".")) if discount_str_input else 0.0
        except ValueError: errors.append("Desconto (%) inválido.")

        if errors:
            messagebox.showerror("Erros de Validação", "\n".join(errors), parent=self)
            return

        # --- Preparação dos Dados para Placeholders ---
        nome_upper=nome.upper(); endereco_final=f"{endereco}, {numero_comp}"; fabricacao_modelo,placa="",""
        if " - " in veiculo:
            try: fabricacao_modelo, placa = veiculo.split(" - ", 1)
            except ValueError: fabricacao_modelo = veiculo
        else: fabricacao_modelo = veiculo

        # Valor diária por extenso
        try:
            valor_diaria_num = float(valor_diaria_str.replace(".","").replace(",","."))
            valor_diaria_extenso = num2words(valor_diaria_num, lang='pt_BR', to='currency').upper()
        except:
            valor_diaria_extenso = f"R$ {valor_diaria_str}"

        # Desconto formatado para o documento
        desconto_perc_val = float(discount_str_input.replace(",",".")) if discount_str_input else 0.0
        desconto_docx_str = f"{desconto_perc_val:.1f}%".replace(".0%", "%").replace(".",",") # Ex: 10% ou 12,5%

        # --- Placeholders ---
        placeholders = {
            "{{NOME}}": nome_upper,
            "{{CPF}}": cpf,
            "{{ENDERECO}}": endereco_final,
            "{{TELEFONE}}": telefone,
            "{{EMAIL}}": email,
            "{{VEICULO}}": veiculo,
            "{{FABRICACAO_MODELO}}": fabricacao_modelo,
            "{{PLACA}}": placa,
            "{{CHASSI}}": chassi,
            "{{RENAVAM}}": renavam,
            "{{PROPRIETARIO}}": proprietario,
            "{{CPF_CNPJ_PROPRIETARIO}}": cpf_cnpj_proprietario,
            "{{PERIODO}}": f"{periodo_str} dias", # Ex: "30 dias"
            "{{PERIODO_EXTENSO}}": periodo_extenso, # Ex: "30 (TRINTA) DIAS"
            "{{VALOR}}": valor_diaria_str, # Valor da diária formatado
            "{{VALOR_EXTENSO}}": valor_diaria_extenso, # Valor da diária por extenso
            "{{VALOR_TOTAL}}": valor_total_str, # Valor total (periodo*diaria) formatado
            "{{VALOR_TOTAL_EXTENSO}}": valor_total_extenso, # Valor total por extenso
            "{{DESCONTO}}": desconto_docx_str, # Desconto formatado Ex: "10%"
            "{{RESULTADO}}": resultado_str, # Valor final com desconto formatado
            "{{DATA}}": data_str
        }

        # --- Seleção e Processamento do DOCX ---
        try:
            model_file = "doc/modelo_contrato_prazodeterminado.docx"
            doc = Document(model_file)
        except FileNotFoundError: return messagebox.showerror("Erro", f"Modelo '{model_file}' não encontrado.", parent=self)
        except Exception as e: return messagebox.showerror("Erro", f"Erro ao abrir modelo:\n{e}", parent=self)

        try:
            # Substituição (mesma lógica robusta da outra frame)
            for para in doc.paragraphs:
                for code, text in placeholders.items():
                    if code in para.text:
                        for run in para.runs:
                            if code in run.text: run.text = run.text.replace(code, str(text))
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                             for code, text in placeholders.items():
                                if code in para.text:
                                    for run in para.runs:
                                         if code in run.text: run.text = run.text.replace(code, str(text))

            # Salvar
            safe_nome = re.sub(r'[^\w\-\.]+', '_', nome_upper).strip('_')
            safe_veiculo = re.sub(r'[^\w\-\.]+', '_', veiculo).strip('_')
            safe_data = data_str.replace("/", "-")
            timestamp = datetime.now().strftime("%H%M%S")
            novo_nome = f"Contrato_Determinado_{safe_nome}_{safe_veiculo}_{safe_data}_{timestamp}.docx"

            doc.save(novo_nome)
            messagebox.showinfo("Sucesso", f"Contrato gerado com sucesso!\n\nArquivo salvo como:\n{novo_nome}", parent=self)

        except Exception as e:
             messagebox.showerror("Erro", f"Erro ao preencher ou salvar o contrato DOCX:\n{e}", parent=self)

class SelecaoFrame(ctk.CTkFrame):
    def __init__(self, master, controller):
        super().__init__(master)
        self.controller = controller
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        self.grid_rowconfigure(2, weight=1)
        self.grid_rowconfigure(3, weight=1)

        label = ctk.CTkLabel(self, text="Selecione a Ação Desejada", font=ctk.CTkFont(size=20, weight="bold"))
        label.grid(row=0, column=0, columnspan=2, padx=20, pady=20, sticky="s")

        btn_determinado = ctk.CTkButton(self, text="Gerar Contrato Prazo Determinado",
                                        command=lambda: controller.show_frame("GeradorDeterminadoFrame"))
        btn_determinado.grid(row=1, column=0, padx=20, pady=10, sticky="ew")

        btn_indeterminado = ctk.CTkButton(self, text="Gerar Contrato Prazo Indeterminado",
                                          command=lambda: controller.show_frame("GeradorIndeterminadoFrame"))
        btn_indeterminado.grid(row=1, column=1, padx=20, pady=10, sticky="ew")

        btn_termo_divida = ctk.CTkButton(self, text="Gerar Termo de Dívida",
                                         command=lambda: controller.show_frame("TermoDividaFrame"))
        btn_termo_divida.grid(row=2, column=0, padx=20, pady=10, sticky="ew")

        btn_investidor = ctk.CTkButton(self, text="Gerar Contrato de Investidor",
                                       command=lambda: controller.show_frame("InvestidorFrame"))
        btn_investidor.grid(row=2, column=1, padx=20, pady=10, sticky="ew")


# --- Aplicação Principal ---
class App(ctk.CTk):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)

        ctk.set_appearance_mode("Light")
        ctk.set_default_color_theme("blue")

        self.title("IZI CAR - Gerador de Contratos") # Título mais limpo
        self.geometry("800x700") # Aumentei um pouco mais

        try:
            # !! IMPORTANTE: Verifique se este caminho está correto !!
            self.iconbitmap("icons/icone.ico")
        except tk.TclError:
            print("Aviso: Ícone 'icons/icone.ico' não encontrado.")
        except Exception as e:
            print(f"Erro ao carregar ícone: {e}")


        # --- Menu Bar ---
        menu_bg_color = "#f0f0f0"
        menu_fg_color = "black"
        menubar = tk.Menu(self, background=menu_bg_color, foreground=menu_fg_color)
        filemenu = tk.Menu(menubar, tearoff=0, background=menu_bg_color, foreground=menu_fg_color)
        filemenu.add_command(label="Sair", command=self.confirm_quit) # Usar confirmação
        menubar.add_cascade(label="Arquivo", menu=filemenu)
        helpmenu = tk.Menu(menubar, tearoff=0, background=menu_bg_color, foreground=menu_fg_color)
        helpmenu.add_command(label="Sobre", command=self.show_about)
        menubar.add_cascade(label="Ajuda", menu=helpmenu)
        self.config(menu=menubar)

        # --- Container Principal ---
        container = ctk.CTkFrame(self, fg_color="transparent") # Fundo transparente
        container.pack(side="top", fill="both", expand=True, padx=10, pady=(10, 0))
        container.grid_rowconfigure(0, weight=1)
        container.grid_columnconfigure(0, weight=1)

        # --- Status Bar ---
        self.statusbar = ctk.CTkLabel(self, text="Status: Pronto", height=25,
                                     anchor="w", font=ctk.CTkFont(size=11))
        self.statusbar.pack(side="bottom", fill="x", padx=10, pady=(5, 5))

        # --- Dicionário de Frames ---
        self.frames = {}

        # --- Instancia os Frames ---
        for F in (SelecaoFrame, GeradorIndeterminadoFrame, GeradorDeterminadoFrame, TermoDividaFrame, InvestidorFrame):
            frame_name = F.__name__
            frame = F(master=container, controller=self)
            self.frames[frame_name] = frame
            frame.grid(row=0, column=0, sticky="nsew")

        # --- Mostra Frame Inicial ---
        self.show_frame("SelecaoFrame")

        # Bind para fechar a janela
        self.protocol("WM_DELETE_WINDOW", self.confirm_quit)


    def show_frame(self, frame_name):
        """ Mostra o frame solicitado """
        if frame_name in self.frames:
            frame = self.frames[frame_name]
            frame.tkraise()
            self.update_statusbar(f"Exibindo: {frame_name.replace('Frame','')}") # Status mais claro
        else:
            print(f"Erro: Frame '{frame_name}' não encontrado.")
            messagebox.showerror("Erro Interno", f"Tentativa de acessar tela inexistente: {frame_name}")

    def update_statusbar(self, message="Pronto"):
        """ Atualiza a barra de status (agora método separado) """
        self.statusbar.configure(text=f"Status: {message}")
        self.update_idletasks() # Força atualização da UI

    def show_about(self):
        """ Mostra a janela 'Sobre' """
        messagebox.showinfo("Sobre IZI CAR Gerador",
                            "Gerador de Contratos IZI CAR\n\nVersão: 1.1 (Integrada)\nDesenvolvido com Python e CustomTkinter.")

    def confirm_quit(self):
        """ Pede confirmação antes de sair """
        if messagebox.askokcancel("Sair", "Tem certeza que deseja sair do aplicativo?", icon='warning'):
            self.quit()
            self.destroy() # Garante que a janela feche completamente


if __name__ == "__main__":
    # Garanta que as dependências estão instaladas:
    # pip install customtkinter requests num2words python-docx
    app = App()
    app.mainloop()
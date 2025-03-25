import tkinter as tk
from tkinter import ttk, messagebox
import re
from datetime import datetime
import requests
import json
from num2words import num2words
from docx import Document
import subprocess  # Importado para executar outro script

try:
    from tkcalendar import DateEntry
except ImportError:
    DateEntry = None

def somente_letras(texto):
    return re.fullmatch(r'[A-Za-zÀ-ÖØ-öø-ÿ\s]+', texto) is not None

def format_cpf(event):
    s = cpf_var.get()
    digits = re.sub(r'\D', '', s)[:11]
    if len(digits) <= 3:
        formatted = digits
    elif len(digits) <= 6:
        formatted = f"{digits[:3]}.{digits[3:]}"
    elif len(digits) <= 9:
        formatted = f"{digits[:3]}.{digits[3:6]}.{digits[6:]}"
    else:
        formatted = f"{digits[:3]}.{digits[3:6]}.{digits[6:9]}-{digits[9:]}"
    cpf_var.set(formatted)

def format_valor(event):
    s = valor_var.get()
    digits = re.sub(r'\D', '', s)
    if digits == "":
        formatted = ""
    else:
        num = int(digits)
        formatted = f"{num:,.2f}"
        formatted = formatted.replace(",", "X").replace(".", ",").replace("X", ".")
    valor_var.set(formatted)

def format_telefone(event):
    s = telefone_var.get()
    digits = re.sub(r'\D', '', s)[:11]
    if len(digits) < 11:
        telefone_var.set(digits)
    else:
        formatted = f"({digits[:2]}){digits[2]} {digits[3:7]}-{digits[7:]}"
        telefone_var.set(formatted)

def buscar_cep():
    cep = cep_var.get().strip().replace('-', '')
    if not re.fullmatch(r'\d{8}', cep):
        messagebox.showerror("Erro", "CEP inválido. Informe 8 dígitos numéricos.")
        return
    try:
        response = requests.get(f"https://viacep.com.br/ws/{cep}/json/")
        data = response.json()
        if data.get("erro"):
            messagebox.showerror("Erro", "CEP não encontrado.")
            return
        logradouro = data.get("logradouro", "")
        bairro = data.get("bairro", "")
        localidade = data.get("localidade", "")
        uf = data.get("uf", "")
        endereco_completo = f"{logradouro}, {bairro}, {localidade} - {uf}"
        endereco_var.set(endereco_completo)
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao buscar CEP: {e}")

try:
    with open("veiculos.json", "r") as f:
        veiculos_data = json.load(f)
except Exception as e:
    messagebox.showerror("Erro", f"Erro ao carregar 'veiculos.json': {e}")
    veiculos_data = {}

veiculo_keys = list(veiculos_data.keys())

def atualizar_dados_veiculo(event):
    key = veiculo_var.get()
    if key in veiculos_data:
        renavam_var.set(veiculos_data[key][0])
        chassi_var.set(veiculos_data[key][1])
        proprietario_var.set(veiculos_data[key][2])
        cpf_cnpj_proprietario_var.set(veiculos_data[key][3])

# Variáveis globais para o parcelamento da caução
qtd_parcelas_caucao = ""
valor_parcela_caucao = 0.0
valor_extenso_caucao = ""

def abrir_configuracao_caucao():
    global qtd_parcelas_caucao, valor_parcela_caucao, valor_extenso_caucao
    top = tk.Toplevel(root)
    top.title("Configurar Parcelamento do Caução")
    top.geometry("350x200")
    
    tk.Label(top, text="Valor de Entrada (R$):").grid(row=0, column=0, padx=10, pady=10, sticky="w")
    entrada_var = tk.StringVar()
    entrada_entry = ttk.Entry(top, textvariable=entrada_var)
    entrada_entry.grid(row=0, column=1, padx=10, pady=10)
    
    tk.Label(top, text="Quantidade de Semanas:").grid(row=1, column=0, padx=10, pady=10, sticky="w")
    semanas_var = tk.StringVar()
    semanas_entry = ttk.Entry(top, textvariable=semanas_var)
    semanas_entry.grid(row=1, column=1, padx=10, pady=10)
    
    def confirmar_parcelamento():
        global qtd_parcelas_caucao, valor_parcela_caucao, valor_extenso_caucao
        try:
            entrada = float(entrada_var.get().replace(",", "."))
            semanas = int(semanas_var.get())
            if entrada < 0 or entrada > 1000:
                messagebox.showerror("Erro", "Valor de entrada deve ser entre 0 e 1000.")
                return
            if semanas < 1:
                messagebox.showerror("Erro", "Quantidade de semanas deve ser no mínimo 1.")
                return
            restante = 1000 - entrada
            parcela = round(restante / semanas, 2)
            qtd_parcelas_caucao = str(semanas)
            valor_parcela_caucao = parcela
            valor_extenso_caucao = num2words(parcela, lang='pt_BR').upper() + " REAIS"
            top.destroy()
        except Exception as e:
            messagebox.showerror("Erro", f"Erro: {e}")
    
    btn_confirma = ttk.Button(top, text="Confirmar", command=confirmar_parcelamento)
    btn_confirma.grid(row=2, column=0, columnspan=2, pady=10)

def on_caucao_change(event):
    if cautao_var.get() == "Parcelado":
        abrir_configuracao_caucao()
    else:
        global qtd_parcelas_caucao, valor_parcela_caucao, valor_extenso_caucao
        qtd_parcelas_caucao = ""
        valor_parcela_caucao = 0.0
        valor_extenso_caucao = ""

def gerar_contrato():
    nome = nome_var.get().strip()
    cpf = cpf_var.get().strip()
    veiculo = veiculo_var.get().strip()
    renavam = renavam_var.get().strip()
    chassi = chassi_var.get().strip()
    proprietario = proprietario_var.get().strip()
    cpf_cnpj_proprietario = cpf_cnpj_proprietario_var.get().strip()
    cep = cep_var.get().strip()
    endereco = endereco_var.get().strip()
    numero = numero_var.get().strip()
    valor = valor_var.get().strip()
    cautao = cautao_var.get().strip()
    data = data_var.get().strip()
    email = email_var.get().strip()
    telefone = telefone_var.get().strip()
    
    # Validações básicas
    if not nome:
        messagebox.showerror("Erro", "O campo Nome é obrigatório.")
        return
    if not somente_letras(nome):
        messagebox.showerror("Erro", "Nome deve conter apenas letras e espaços.")
        return
    if not re.fullmatch(r'\d{3}\.\d{3}\.\d{3}-\d{2}', cpf):
        messagebox.showerror("Erro", "CPF deve estar no formato 000.000.000-00.")
        return
    if not veiculo or veiculo == "Selecione...":
        messagebox.showerror("Erro", "Selecione um veículo.")
        return
    if not renavam:
        messagebox.showerror("Erro", "Renavam não preenchido.")
        return
    if not chassi:
        messagebox.showerror("Erro", "Chassi não preenchido.")
        return
    if not proprietario:
        messagebox.showerror("Erro", "Proprietário não preenchido.")
        return
    if not cpf_cnpj_proprietario:
        messagebox.showerror("Erro", "CPF/CNPJ do Proprietário não preenchido.")
        return
    if not cep:
        messagebox.showerror("Erro", "O campo CEP é obrigatório.")
        return
    if not endereco:
        messagebox.showerror("Erro", "Não foi possível obter o endereço a partir do CEP.")
        return
    if not numero or not re.fullmatch(r'\d+', numero):
        messagebox.showerror("Erro", "Número da residência é obrigatório e deve conter apenas dígitos.")
        return
    if not re.fullmatch(r'(\d{1,3}(?:\.\d{3})*|\d+),\d{2}', valor):
        messagebox.showerror("Erro", "Valor deve estar no formato de contabilidade (ex: 100,00 ou 1.000,00).")
        return
    if not re.fullmatch(r'^[\w\.-]+@[\w\.-]+\.\w+$', email):
        messagebox.showerror("Erro", "Email inválido.")
        return
    digits_tel = re.sub(r'\D', '', telefone)
    if len(digits_tel) != 11:
        messagebox.showerror("Erro", "Telefone inválido. Informe 11 dígitos.")
        return
    try:
        datetime.strptime(data, "%d/%m/%Y")
    except ValueError:
        messagebox.showerror("Erro", "Data inválida. Use o formato dd/mm/aaaa.")
        return

    nome = nome.upper()
    
    if " - " in veiculo:
        fabricacao_modelo, placa = veiculo.split(" - ", 1)
    else:
        fabricacao_modelo, placa = veiculo, ""
    
    try:
        valor_num = float(valor.replace(".", "").replace(",", "."))
    except:
        valor_num = 0
    valor_extenso = num2words(valor_num, lang='pt_BR').upper() + " REAIS"
    
    endereco_final = f"{numero}, {endereco}" if numero and endereco else endereco
    
    if cautao == "Parcelado":
        if not qtd_parcelas_caucao:
            qtd = "50"
            valor_parc = 20.00
            valor_parc_fmt = f"{valor_parc:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            valor_extenso_caucao_local = num2words(valor_parc, lang='pt_BR').upper() + " REAIS"
        else:
            qtd = qtd_parcelas_caucao
            valor_parc_fmt = f"{valor_parcela_caucao:,.2f}"
            valor_parc_fmt = valor_parc_fmt.replace(",", "X").replace(".", ",").replace("X", ".")
            valor_extenso_caucao_local = valor_extenso_caucao
    else:
        qtd = ""
        valor_parc_fmt = ""
        valor_extenso_caucao_local = ""
    
    if cautao == "Parcelado":
        model_file = "modelo_contrato_prazoinderteminadoparcelado.docx"
    else:
        model_file = "modelo_contrato_prazoinderteminadoquitado.docx"
    
    placeholders = {
        "{{NOME}}": nome,
        "{{CPF}}": cpf,
        "{{ENDERECO}}": endereco_final,
        "{{VEICULO}}": veiculo,
        "{{FABRICACAO_MODELO}}": fabricacao_modelo,
        "{{PLACA}}": placa,
        "{{CHASSI}}": chassi,
        "{{RENAVAM}}": renavam,
        "{{PROPRIETARIO}}": proprietario,
        "{{CPF_CNPJ_PROPRIETARIO}}": cpf_cnpj_proprietario,
        "{{VALOR}}": valor,
        "{{VALOR_EXTENSO}}": valor_extenso,
        "{{QTD_PARCELAS}}": qtd,
        "{{VALOR_PARCELAS_CAUCAO}}": valor_parc_fmt,
        "{{VALOR_EXTENSO_CAUCAO}}": valor_extenso_caucao_local,
        "{{DATA}}": data,
        "{{EMAIL}}": email,
        "{{TELEFONE}}": telefone
    }
    
    try:
        doc = Document(model_file)
        for para in doc.paragraphs:
            for ph, val in placeholders.items():
                if ph in para.text:
                    para.text = para.text.replace(ph, val)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for ph, val in placeholders.items():
                        if ph in cell.text:
                            cell.text = cell.text.replace(ph, val)
        safe_nome = nome.replace(" ", "_")
        safe_veiculo = veiculo.replace(" ", "_")
        safe_data = data.replace("/", "-")
        novo_nome = f"Contrato_{safe_nome}_{safe_veiculo}_{safe_data}.docx"
        doc.save(novo_nome)
        messagebox.showinfo("Sucesso", f"Contrato gerado com sucesso: {novo_nome}")
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao gerar o contrato: {e}")

# Função para voltar: exibe mensagem de confirmação, fecha a janela atual e executa "PAINELENTRADA.py"
def voltar():
    if messagebox.askokcancel("Confirmar", "Tem certeza que quer voltar? Todo o progresso será perdido"):
        root.destroy()
        subprocess.Popen(["python", "PAINELENTRADA.py"])

root = tk.Tk()
root.iconbitmap("icone.ico")
root.title("Gerador de Contratos")
root.geometry("600x750")
root.configure(bg="white")

nome_var = tk.StringVar()
cpf_var = tk.StringVar()
veiculo_var = tk.StringVar(value="Selecione...")
renavam_var = tk.StringVar()
chassi_var = tk.StringVar()
proprietario_var = tk.StringVar()
cpf_cnpj_proprietario_var = tk.StringVar()
cep_var = tk.StringVar()
endereco_var = tk.StringVar()
numero_var = tk.StringVar()
valor_var = tk.StringVar()
cautao_var = tk.StringVar(value="Quitado")
data_var = tk.StringVar()
email_var = tk.StringVar()
telefone_var = tk.StringVar()

main_frame = ttk.Frame(root, padding=20)
main_frame.pack(fill="both", expand=True)

header = ttk.Label(main_frame, text="GERADOR DE CONTRATO IZI CAR", style="Header.TLabel")
header.grid(row=0, column=0, columnspan=3, pady=(0,20))

ttk.Label(main_frame, text="Nome:").grid(row=1, column=0, sticky="w", pady=5)
nome_entry = ttk.Entry(main_frame, textvariable=nome_var)
nome_entry.grid(row=1, column=1, columnspan=2, sticky="ew", pady=5)

ttk.Label(main_frame, text="CPF:").grid(row=2, column=0, sticky="w", pady=5)
cpf_entry = ttk.Entry(main_frame, textvariable=cpf_var)
cpf_entry.grid(row=2, column=1, columnspan=2, sticky="ew", pady=5)
cpf_entry.bind("<FocusOut>", format_cpf)

ttk.Label(main_frame, text="Veículo:").grid(row=3, column=0, sticky="w", pady=5)
veiculo_cb = ttk.Combobox(main_frame, textvariable=veiculo_var, values=veiculo_keys, state="readonly")
veiculo_cb.grid(row=3, column=1, columnspan=2, sticky="ew", pady=5)
veiculo_cb.set("Selecione...")
veiculo_cb.bind("<<ComboboxSelected>>", atualizar_dados_veiculo)

ttk.Label(main_frame, text="Renavam:").grid(row=4, column=0, sticky="w", pady=5)
renavam_entry = ttk.Entry(main_frame, textvariable=renavam_var, state="readonly")
renavam_entry.grid(row=4, column=1, columnspan=2, sticky="ew", pady=5)

ttk.Label(main_frame, text="Chassi:").grid(row=5, column=0, sticky="w", pady=5)
chassi_entry = ttk.Entry(main_frame, textvariable=chassi_var, state="readonly")
chassi_entry.grid(row=5, column=1, columnspan=2, sticky="ew", pady=5)

ttk.Label(main_frame, text="Proprietário:").grid(row=6, column=0, sticky="w", pady=5)
proprietario_entry = ttk.Entry(main_frame, textvariable=proprietario_var, state="readonly")
proprietario_entry.grid(row=6, column=1, columnspan=2, sticky="ew", pady=5)

ttk.Label(main_frame, text="CPF/CNPJ Proprietário:").grid(row=7, column=0, sticky="w", pady=5)
cpf_cnpj_entry = ttk.Entry(main_frame, textvariable=cpf_cnpj_proprietario_var, state="readonly")
cpf_cnpj_entry.grid(row=7, column=1, columnspan=2, sticky="ew", pady=5)

ttk.Label(main_frame, text="CEP:").grid(row=8, column=0, sticky="w", pady=5)
cep_entry = ttk.Entry(main_frame, textvariable=cep_var)
cep_entry.grid(row=8, column=1, sticky="ew", pady=5)
btn_cep = ttk.Button(main_frame, text="Buscar CEP", command=buscar_cep)
btn_cep.grid(row=8, column=2, padx=5, pady=5)

ttk.Label(main_frame, text="Endereço:").grid(row=9, column=0, sticky="w", pady=5)
endereco_entry = ttk.Entry(main_frame, textvariable=endereco_var, state="readonly")
endereco_entry.grid(row=9, column=1, columnspan=2, sticky="ew", pady=5)

ttk.Label(main_frame, text="Número/Complemento:").grid(row=10, column=0, sticky="w", pady=5)
numero_entry = ttk.Entry(main_frame, textvariable=numero_var)
numero_entry.grid(row=10, column=1, columnspan=2, sticky="ew", pady=5)

ttk.Label(main_frame, text="Email:").grid(row=11, column=0, sticky="w", pady=5)
email_entry = ttk.Entry(main_frame, textvariable=email_var)
email_entry.grid(row=11, column=1, columnspan=2, sticky="ew", pady=5)

ttk.Label(main_frame, text="Telefone:").grid(row=12, column=0, sticky="w", pady=5)
telefone_entry = ttk.Entry(main_frame, textvariable=telefone_var)
telefone_entry.grid(row=12, column=1, columnspan=2, sticky="ew", pady=5)
telefone_entry.bind("<FocusOut>", format_telefone)

ttk.Label(main_frame, text="Valor:").grid(row=13, column=0, sticky="w", pady=5)
valor_entry = ttk.Entry(main_frame, textvariable=valor_var)
valor_entry.grid(row=13, column=1, columnspan=2, sticky="ew", pady=5)
valor_entry.bind("<FocusOut>", format_valor)

ttk.Label(main_frame, text="Caução:").grid(row=14, column=0, sticky="w", pady=5)
cautao_cb = ttk.Combobox(main_frame, textvariable=cautao_var, values=["Quitado", "Parcelado"], state="readonly")
cautao_cb.grid(row=14, column=1, columnspan=2, sticky="ew", pady=5)
cautao_cb.bind("<<ComboboxSelected>>", on_caucao_change)

ttk.Label(main_frame, text="Data (dd/mm/aaaa):").grid(row=15, column=0, sticky="w", pady=5)
if DateEntry:
    data_entry = DateEntry(main_frame, textvariable=data_var, date_pattern="dd/mm/yyyy")
else:
    data_entry = ttk.Entry(main_frame, textvariable=data_var)
data_entry.grid(row=15, column=1, columnspan=2, sticky="ew", pady=5)

btn_gerar = ttk.Button(main_frame, text="GERAR CONTRATO", command=gerar_contrato)
btn_gerar.grid(row=16, column=0, columnspan=3, pady=20)

# Botão Voltar: fecha a janela atual e executa o script "PAINELENTRADA.py" após confirmação
btn_voltar = ttk.Button(main_frame, text="Voltar", command=voltar)
btn_voltar.grid(row=17, column=0, columnspan=3, pady=10)

main_frame.columnconfigure(0, weight=1)
main_frame.columnconfigure(1, weight=3)
main_frame.columnconfigure(2, weight=1)

root.mainloop()

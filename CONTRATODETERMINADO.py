import tkinter as tk
from tkinter import ttk, messagebox
import re
from datetime import datetime
import requests
import json
from num2words import num2words
from docx import Document
import subprocess

try:
    from tkcalendar import DateEntry
except ImportError:
    DateEntry = None

# Funções de validação e formatação

def somente_letras(texto):
    return re.fullmatch(r'[A-Za-zÀ-ÖØ-öø-ÿ\s]+', texto) is not None

def format_cpf(event):
    s = cpf_var.get()
    digits = re.sub(r'\D', '', s)[:11]
    if len(digits) <= 3:
        formatted = digits
    elif len(digits) <= 6:
        formatted = f"{digits[:3]}.{digits[3:]}"
    elif len(digits) <= 9:
        formatted = f"{digits[:3]}.{digits[3:6]}.{digits[6:]}"
    else:
        formatted = f"{digits[:3]}.{digits[3:6]}.{digits[6:9]}-{digits[9:]}"
    cpf_var.set(formatted)

def format_valor(event):
    s = valor_var.get()
    digits = re.sub(r'\D', '', s)
    if digits == "":
        formatted = ""
    else:
        num = int(digits)
        formatted = f"{num:,.2f}"
        formatted = formatted.replace(",", "X").replace(".", ",").replace("X", ".")
    valor_var.set(formatted)

def format_valor_total(event):
    s = valor_total_var.get()
    digits = re.sub(r'\D', '', s)
    if digits == "":
        formatted = ""
    else:
        num = int(digits)
        formatted = f"{num:,.2f}"
        formatted = formatted.replace(",", "X").replace(".", ",").replace("X", ".")
    valor_total_var.set(formatted)
    try:
        valor_num = float(digits)
    except:
        valor_num = 0
    valor_total_extenso_var.set(num2words(valor_num, lang='pt_BR').upper() + " REAIS")

def format_periodo(event):
    s = periodo_var.get().strip()
    if s.isdigit():
        periodo_extenso_var.set(f"{s} dias - {num2words(int(s), lang='pt_BR').upper()} DIAS")
    else:
        periodo_extenso_var.set("")

def format_telefone(event):
    s = telefone_var.get()
    digits = re.sub(r'\D', '', s)[:11]
    if len(digits) < 11:
        telefone_var.set(digits)
    else:
        formatted = f"({digits[:2]}){digits[2]} {digits[3:7]}-{digits[7:]}"
        telefone_var.set(formatted)

def buscar_cep():
    cep = cep_var.get().strip().replace('-', '')
    if not re.fullmatch(r'\d{8}', cep):
        messagebox.showerror("Erro", "CEP inválido. Informe 8 dígitos numéricos.")
        return
    try:
        response = requests.get(f"https://viacep.com.br/ws/{cep}/json/")
        data = response.json()
        if data.get("erro"):
            messagebox.showerror("Erro", "CEP não encontrado.")
            return
        logradouro = data.get("logradouro", "")
        bairro = data.get("bairro", "")
        localidade = data.get("localidade", "")
        uf = data.get("uf", "")
        endereco_completo = f"{logradouro}, {bairro}, {localidade} - {uf}"
        endereco_var.set(endereco_completo)
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao buscar CEP: {e}")

# Carregar dados dos veículos do JSON
try:
    with open("veiculos.json", "r") as f:
        veiculos_data = json.load(f)
except Exception as e:
    messagebox.showerror("Erro", f"Erro ao carregar 'veiculos.json': {e}")
    veiculos_data = {}

veiculo_keys = list(veiculos_data.keys())

def atualizar_dados_veiculo(event):
    key = veiculo_var.get()
    if key in veiculos_data:
        renavam_var.set(veiculos_data[key][0])
        chassi_var.set(veiculos_data[key][1])
        proprietario_var.set(veiculos_data[key][2])
        cpf_cnpj_proprietario_var.set(veiculos_data[key][3])

# Variáveis globais para o parcelamento da caução (não utilizadas nesta interface)
qtd_parcelas_caucao = ""
valor_parcela_caucao = 0.0
valor_extenso_caucao = ""

def abrir_configuracao_caucao():
    global qtd_parcelas_caucao, valor_parcela_caucao, valor_extenso_caucao
    top = tk.Toplevel(root)
    top.title("Configurar Parcelamento do Caução")
    top.geometry("350x200")
    
    tk.Label(top, text="Valor de Entrada (R$):").grid(row=0, column=0, padx=10, pady=10, sticky="w")
    entrada_var = tk.StringVar()
    entrada_entry = ttk.Entry(top, textvariable=entrada_var)
    entrada_entry.grid(row=0, column=1, padx=10, pady=10)
    
    tk.Label(top, text="Quantidade de Semanas:").grid(row=1, column=0, padx=10, pady=10, sticky="w")
    semanas_var = tk.StringVar()
    semanas_entry = ttk.Entry(top, textvariable=semanas_var)
    semanas_entry.grid(row=1, column=1, padx=10, pady=10)
    
    def confirmar_parcelamento():
        global qtd_parcelas_caucao, valor_parcela_caucao, valor_extenso_caucao
        try:
            entrada = float(entrada_var.get().replace(",", "."))
            semanas = int(semanas_var.get())
            if entrada < 0 or entrada > 1000:
                messagebox.showerror("Erro", "Valor de entrada deve ser entre 0 e 1000.")
                return
            if semanas < 1:
                messagebox.showerror("Erro", "Quantidade de semanas deve ser no mínimo 1.")
                return
            restante = 1000 - entrada
            parcela = round(restante / semanas, 2)
            qtd_parcelas_caucao = str(semanas)
            valor_parcela_caucao = parcela
            valor_extenso_caucao = num2words(parcela, lang='pt_BR').upper() + " REAIS"
            top.destroy()
        except Exception as e:
            messagebox.showerror("Erro", f"Erro: {e}")
    
    btn_confirma = ttk.Button(top, text="Confirmar", command=confirmar_parcelamento)
    btn_confirma.grid(row=2, column=0, columnspan=2, pady=10)

def on_caucao_change(event):
    if cautao_var.get() == "Parcelado":
        abrir_configuracao_caucao()
    else:
        global qtd_parcelas_caucao, valor_parcela_caucao, valor_extenso_caucao
        qtd_parcelas_caucao = ""
        valor_parcela_caucao = 0.0
        valor_extenso_caucao = ""

def gerar_contrato():
    nome = nome_var.get().strip()
    cpf = cpf_var.get().strip()
    veiculo = veiculo_var.get().strip()
    renavam = renavam_var.get().strip()
    chassi = chassi_var.get().strip()
    proprietario = proprietario_var.get().strip()
    cpf_cnpj_proprietario = cpf_cnpj_proprietario_var.get().strip()
    cep = cep_var.get().strip()
    endereco = endereco_var.get().strip()
    numero = numero_var.get().strip()
    valor = valor_var.get().strip()
    cautao = cautao_var.get().strip()
    data = data_var.get().strip()
    email = email_var.get().strip()
    telefone = telefone_var.get().strip()
    periodo = periodo_var.get().strip()
    
    if not nome:
        messagebox.showerror("Erro", "O campo Nome é obrigatório.")
        return
    if not somente_letras(nome):
        messagebox.showerror("Erro", "Nome deve conter apenas letras e espaços.")
        return
    if not re.fullmatch(r'\d{3}\.\d{3}\.\d{3}-\d{2}', cpf):
        messagebox.showerror("Erro", "CPF deve estar no formato 000.000.000-00.")
        return
    if not veiculo or veiculo == "Selecione...":
        messagebox.showerror("Erro", "Selecione um veículo.")
        return
    if not renavam:
        messagebox.showerror("Erro", "Renavam não preenchido.")
        return
    if not chassi:
        messagebox.showerror("Erro", "Chassi não preenchido.")
        return
    if not proprietario:
        messagebox.showerror("Erro", "Proprietário não preenchido.")
        return
    if not cpf_cnpj_proprietario:
        messagebox.showerror("Erro", "CPF/CNPJ do Proprietário não preenchido.")
        return
    if not cep:
        messagebox.showerror("Erro", "O campo CEP é obrigatório.")
        return
    if not endereco:
        messagebox.showerror("Erro", "Não foi possível obter o endereço a partir do CEP.")
        return
    if not numero or not re.fullmatch(r'\d+', numero):
        messagebox.showerror("Erro", "Número da residência é obrigatório e deve conter apenas dígitos.")
        return
    if not re.fullmatch(r'(\d{1,3}(?:\.\d{3})*|\d+),\d{2}', valor):
        messagebox.showerror("Erro", "Valor deve estar no formato de contabilidade (ex: 100,00 ou 1.000,00).")
        return
    if not re.fullmatch(r'^[\w\.-]+@[\w\.-]+\.\w+$', email):
        messagebox.showerror("Erro", "Email inválido.")
        return
    digits_tel = re.sub(r'\D', '', telefone)
    if len(digits_tel) != 11:
        messagebox.showerror("Erro", "Telefone inválido. Informe 11 dígitos.")
        return
    try:
        datetime.strptime(data, "%d/%m/%Y")
    except ValueError:
        messagebox.showerror("Erro", "Data inválida. Use o formato dd/mm/aaaa.")
        return
    if not periodo or not periodo.isdigit():
        messagebox.showerror("Erro", "O campo PERIODO é obrigatório e deve conter apenas números.")
        return

    nome = nome.upper()
    
    if " - " in veiculo:
        fabricacao_modelo, placa = veiculo.split(" - ", 1)
    else:
        fabricacao_modelo, placa = veiculo, ""
    
    try:
        valor_num = float(valor.replace(".", "").replace(",", "."))
    except:
        valor_num = 0
    valor_extenso = num2words(valor_num, lang='pt_BR').upper() + " REAIS"
    
    endereco_final = f"{numero}, {endereco}"
    
    period_int = int(periodo)
    total = valor_num * period_int
    total_fmt = f"{total:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    total_extenso = num2words(total, lang='pt_BR').upper() + " REAIS"
    periodo_str = f"{period_int} dias"
    periodo_extenso = f"{period_int} dias - {num2words(period_int, lang='pt_BR').upper()} DIAS"
    daily_extenso = num2words(valor_num, lang='pt_BR').upper() + " REAIS"
    
    desconto_val = float(discount_var.get()) if discount_var.get() else 0
    desconto_str = f"{desconto_val}%"
    resultado = total * (1 - desconto_val / 100)
    resultado_fmt = f"{resultado:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    
    placeholders = {
        "{{NOME}}": nome,
        "{{CPF}}": cpf,
        "{{ENDERECO}}": endereco_final,
        "{{TELEFONE}}": telefone,
        "{{EMAIL}}": email,
        "{{VEICULO}}": veiculo,
        "{{FABRICACAO_MODELO}}": fabricacao_modelo,
        "{{PLACA}}": placa,
        "{{CHASSI}}": chassi,
        "{{RENAVAM}}": renavam,
        "{{PROPRIETARIO}}": proprietario,
        "{{CPF_CNPJ_PROPRIETARIO}}": cpf_cnpj_proprietario,
        "{{PERIODO}}": periodo_str,
        "{{PERIODO_EXTENSO}}": periodo_extenso,
        "{{VALOR}}": valor,
        "{{VALOR_EXTENSO}}": valor_extenso,
        "{{VALOR_TOTAL}}": total_fmt,
        "{{VALOR_TOTAL_EXTENSO}}": total_extenso,
        "{{DESCONTO}}": desconto_str,
        "{{RESULTADO}}": resultado_fmt,
        "{{DATA}}": data
    }
    
    try:
        doc = Document("modelo_contrato_prazodeterminado.docx")
        for para in doc.paragraphs:
            for ph, val in placeholders.items():
                if ph in para.text:
                    para.text = para.text.replace(ph, val)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for ph, val in placeholders.items():
                        if ph in cell.text:
                            cell.text = cell.text.replace(ph, val)
        safe_nome = nome.replace(" ", "_")
        safe_veiculo = veiculo.replace(" ", "_")
        safe_data = data.replace("/", "-")
        novo_nome = f"Contrato_{safe_nome}_{safe_veiculo}_{safe_data}.docx"
        doc.save(novo_nome)
        messagebox.showinfo("Sucesso", f"Contrato gerado com sucesso: {novo_nome}")
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao gerar o contrato: {e}")

root = tk.Tk()
root.iconbitmap("icone.ico")
root.title("Gerador de Contratos")
root.geometry("700x650")
root.configure(bg="light gray")

# Botão para voltar ao Painel Entrada com confirmação
def voltar_painel():
    if messagebox.askyesno("Confirmar", "Tem certeza que quer voltar? Todo o progresso será perdido."):
        subprocess.Popen(["python", "PAINELENTRADA.py"])
        root.destroy()

# Cria um frame para os botões na parte inferior
button_frame = ttk.Frame(root, padding=10)
button_frame.pack(side="bottom", fill="x")

btn_gerar = ttk.Button(button_frame, text="GERAR CONTRATO", command=gerar_contrato)
btn_gerar.pack(side="left", padx=10, pady=10)

btn_voltar = ttk.Button(button_frame, text="Voltar", command=voltar_painel)
btn_voltar.pack(side="left", padx=10, pady=10)

# Variáveis de controle
nome_var = tk.StringVar()
cpf_var = tk.StringVar()
veiculo_var = tk.StringVar(value="Selecione...")
renavam_var = tk.StringVar()
chassi_var = tk.StringVar()
proprietario_var = tk.StringVar()
cpf_cnpj_proprietario_var = tk.StringVar()
cep_var = tk.StringVar()
endereco_var = tk.StringVar()
numero_var = tk.StringVar()
valor_var = tk.StringVar()
cautao_var = tk.StringVar(value="Quitado")
data_var = tk.StringVar()
email_var = tk.StringVar()
telefone_var = tk.StringVar()
periodo_var = tk.StringVar()
discount_var = tk.StringVar()
valor_total_var = tk.StringVar()

valor_total_extenso_var = tk.StringVar()
periodo_extenso_var = tk.StringVar()
result_var = tk.StringVar()

main_frame = ttk.Frame(root, padding=20)
main_frame.pack(fill="both", expand=True)

row = 0
ttk.Label(main_frame, text="NOME:").grid(row=row, column=0, sticky="w", pady=5)
nome_entry = ttk.Entry(main_frame, textvariable=nome_var)
nome_entry.grid(row=row, column=1, columnspan=2, sticky="ew", pady=5)
row += 1

ttk.Label(main_frame, text="CPF:").grid(row=row, column=0, sticky="w", pady=5)
cpf_entry = ttk.Entry(main_frame, textvariable=cpf_var)
cpf_entry.grid(row=row, column=1, columnspan=2, sticky="ew", pady=5)
cpf_entry.bind("<FocusOut>", format_cpf)
row += 1

ttk.Label(main_frame, text="CEP:").grid(row=row, column=0, sticky="w", pady=5)
cep_entry = ttk.Entry(main_frame, textvariable=cep_var)
cep_entry.grid(row=row, column=1, sticky="ew", pady=5)
btn_cep = ttk.Button(main_frame, text="Buscar CEP", command=buscar_cep)
btn_cep.grid(row=row, column=2, padx=5, pady=5)
row += 1

ttk.Label(main_frame, text="ENDEREÇO:").grid(row=row, column=0, sticky="w", pady=5)
endereco_entry = ttk.Entry(main_frame, textvariable=endereco_var, state="readonly")
endereco_entry.grid(row=row, column=1, columnspan=2, sticky="ew", pady=5)
row += 1

ttk.Label(main_frame, text="NÚMERO/COMPLEMENTO:").grid(row=row, column=0, sticky="w", pady=5)
numero_entry = ttk.Entry(main_frame, textvariable=numero_var)
numero_entry.grid(row=row, column=1, columnspan=2, sticky="ew", pady=5)
row += 1

ttk.Label(main_frame, text="EMAIL:").grid(row=row, column=0, sticky="w", pady=5)
email_entry = ttk.Entry(main_frame, textvariable=email_var)
email_entry.grid(row=row, column=1, columnspan=2, sticky="ew", pady=5)
row += 1

ttk.Label(main_frame, text="TELEFONE:").grid(row=row, column=0, sticky="w", pady=5)
telefone_entry = ttk.Entry(main_frame, textvariable=telefone_var)
telefone_entry.grid(row=row, column=1, columnspan=2, sticky="ew", pady=5)
telefone_entry.bind("<FocusOut>", format_telefone)
row += 1

ttk.Label(main_frame, text="VEICULO:").grid(row=row, column=0, sticky="w", pady=5)
veiculo_cb = ttk.Combobox(main_frame, textvariable=veiculo_var, values=veiculo_keys, state="readonly")
veiculo_cb.grid(row=row, column=1, columnspan=2, sticky="ew", pady=5)
veiculo_cb.set("Selecione...")
veiculo_cb.bind("<<ComboboxSelected>>", atualizar_dados_veiculo)
row += 1

ttk.Label(main_frame, text="CHASSI:").grid(row=row, column=0, sticky="w", pady=5)
chassi_entry = ttk.Entry(main_frame, textvariable=chassi_var, state="readonly")
chassi_entry.grid(row=row, column=1, columnspan=2, sticky="ew", pady=5)
row += 1

ttk.Label(main_frame, text="RENAVAM:").grid(row=row, column=0, sticky="w", pady=5)
renavam_entry = ttk.Entry(main_frame, textvariable=renavam_var, state="readonly")
renavam_entry.grid(row=row, column=1, columnspan=2, sticky="ew", pady=5)
row += 1

ttk.Label(main_frame, text="PROPRIETÁRIO:").grid(row=row, column=0, sticky="w", pady=5)
proprietario_entry = ttk.Entry(main_frame, textvariable=proprietario_var, state="readonly")
proprietario_entry.grid(row=row, column=1, columnspan=2, sticky="ew", pady=5)
row += 1

ttk.Label(main_frame, text="CPF/CNPJ PROPRIETÁRIO:").grid(row=row, column=0, sticky="w", pady=5)
cpf_cnpj_entry = ttk.Entry(main_frame, textvariable=cpf_cnpj_proprietario_var, state="readonly")
cpf_cnpj_entry.grid(row=row, column=1, columnspan=2, sticky="ew", pady=5)
row += 1

ttk.Label(main_frame, text="PERIODO (dias):").grid(row=row, column=0, sticky="w", pady=5)
periodo_entry = ttk.Entry(main_frame, textvariable=periodo_var)
periodo_entry.grid(row=row, column=1, columnspan=2, sticky="ew", pady=5)
periodo_entry.bind("<FocusOut>", format_periodo)
row += 1

ttk.Label(main_frame, text="VALOR (diária):").grid(row=row, column=0, sticky="w", pady=5)
valor_entry = ttk.Entry(main_frame, textvariable=valor_var)
valor_entry.grid(row=row, column=1, columnspan=2, sticky="ew", pady=5)
valor_entry.bind("<FocusOut>", format_valor)
row += 1

ttk.Label(main_frame, text="DESCONTO (%):").grid(row=row, column=0, sticky="w", pady=5)
discount_entry = ttk.Entry(main_frame, textvariable=discount_var)
discount_entry.grid(row=row, column=1, columnspan=2, sticky="ew", pady=5)
row += 1

ttk.Label(main_frame, text="RESULTADO:").grid(row=row, column=0, sticky="w", pady=5)
result_entry = ttk.Entry(main_frame, textvariable=result_var, state="readonly")
result_entry.grid(row=row, column=1, columnspan=2, sticky="ew", pady=5)
row += 1

ttk.Label(main_frame, text="DATA (dd/mm/aaaa):").grid(row=row, column=0, sticky="w", pady=5)
if DateEntry:
    data_entry = DateEntry(main_frame, textvariable=data_var, date_pattern="dd/mm/yyyy")
else:
    data_entry = ttk.Entry(main_frame, textvariable=data_var)
data_entry.grid(row=row, column=1, columnspan=2, sticky="ew", pady=5)
row += 1

root.mainloop()
